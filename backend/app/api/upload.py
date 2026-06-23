"""
Media upload & transcription API — multi-tenant edition with graceful fallback.
When Sarvam + Gemini + Pinecone available: full transcription + embedding pipeline.
When unavailable: stores files and provides fallback transcription.
"""

import os
import uuid
import glob
import json
import subprocess
from datetime import datetime, timezone

from fastapi import APIRouter, File, UploadFile, Form, BackgroundTasks, Depends, HTTPException
from pydantic import BaseModel
from typing import Optional
from dotenv import load_dotenv
from sqlalchemy.orm import Session

from ..models import Document, Subscription, User, generate_uuid
from .deps import get_db, get_current_user, get_optional_user
from ..config import SARVAM_API_KEY, GEMINI_API_KEY, PINECONE_API_KEY, PINECONE_INDEX_NAME

router = APIRouter()
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# External service initialization (graceful)
# ---------------------------------------------------------------------------

gemini_client = None
if GEMINI_API_KEY:
    try:
        from google import genai
        gemini_client = genai.Client(api_key=GEMINI_API_KEY)
        print("[OK] Upload: Gemini embedding client initialized.")
    except Exception as e:
        print(f"[WARN] Upload: Gemini client init failed: {e}")

pinecone_index = None
if PINECONE_API_KEY:
    try:
        from pinecone import Pinecone
        pc = Pinecone(api_key=PINECONE_API_KEY)
        pinecone_index = pc.Index(PINECONE_INDEX_NAME)
        print(f"[OK] Upload: Pinecone index '{PINECONE_INDEX_NAME}' connected.")
    except Exception as e:
        print(f"[WARN] Upload: Pinecone connection failed: {e}")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

class UploadResponse(BaseModel):
    file_id: str
    filename: str
    message: str


def chunk_text(text, chunk_size=500):
    words = text.split()
    chunks = []
    current_chunk = []
    current_len = 0
    for word in words:
        if current_len + len(word) > chunk_size:
            chunks.append(" ".join(current_chunk))
            current_chunk = [word]
            current_len = len(word)
        else:
            current_chunk.append(word)
            current_len += len(word) + 1
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    return chunks


# ---------------------------------------------------------------------------
# Transcription: Real pipeline (Sarvam API)
# ---------------------------------------------------------------------------

def transcribe_with_sarvam(file_path: str, file_id: str) -> str:
    """Transcribe audio using Sarvam API with chunked processing."""
    import requests
    import imageio_ffmpeg

    ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
    print(f"[{file_id}] Slicing media into 30s chunks using ffmpeg...")

    chunk_pattern = f"{file_path}_chunk_%03d.mp3"
    subprocess.run([
        ffmpeg_exe, "-y", "-i", file_path,
        "-f", "segment", "-segment_time", "30",
        "-c:a", "libmp3lame", "-q:a", "0", "-map", "a",
        chunk_pattern
    ], check=True)

    full_transcription = ""
    url = "https://api.sarvam.ai/speech-to-text"
    headers = {"api-subscription-key": SARVAM_API_KEY}

    chunk_files = sorted(glob.glob(f"{file_path}_chunk_*.mp3"))

    for idx, chunk_path in enumerate(chunk_files):
        print(f"[{file_id}] Transcribing chunk {idx+1}/{len(chunk_files)}...")
        files = {"file": ("audio.mp3", open(chunk_path, "rb"), "audio/mpeg")}
        data = {"model": "saaras:v3", "mode": "transcribe"}

        response = requests.post(url, headers=headers, files=files, data=data)

        if response.status_code == 200:
            transcript = response.json().get("transcript", "")
            full_transcription += transcript + " "
        else:
            raise Exception(f"Sarvam API Error on chunk: {response.text}")

    # Clean up chunk files
    for chunk_path in chunk_files:
        try:
            os.remove(chunk_path)
        except Exception:
            pass

    return full_transcription.strip()


# ---------------------------------------------------------------------------
# Transcription: Fallback mode (no Sarvam API)
# ---------------------------------------------------------------------------

def transcribe_fallback(file_path: str, file_id: str, original_filename: str) -> str:
    ext = os.path.splitext(original_filename)[1].lower()
    file_size_mb = os.path.getsize(file_path) / (1024 * 1024)

    # PDF files — extract text using pypdf
    if ext in [".pdf"]:
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text.strip() if text.strip() else f"Empty PDF file: {original_filename}"
        except Exception as pdf_err:
            print(f"PDF extraction error: {pdf_err}")
            return f"Error extracting text from PDF file: {original_filename}\n{str(pdf_err)}"

    # Word Documents — extract text using python-docx
    if ext in [".docx"]:
        try:
            import docx
            doc = docx.Document(file_path)
            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
            text = "\n".join(fullText)
            return text.strip() if text.strip() else f"Empty Document: {original_filename}"
        except Exception as docx_err:
            print(f"DOCX extraction error: {docx_err}")
            return f"Error extracting text from DOCX file: {original_filename}\n{str(docx_err)}"

    # Text files — read content directly
    if ext in [".txt", ".text"]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
            return content.strip() if content.strip() else f"Empty text file: {original_filename}"
        except Exception:
            pass

    # CSV files — read as text
    if ext in [".csv"]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
            return f"CSV Data from {original_filename}:\n{content[:10000]}"
        except Exception:
            pass

    # Audio/Video files — provide metadata-based description
    duration_info = ""
    try:
        import imageio_ffmpeg
        ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
        result = subprocess.run(
            [ffmpeg_exe, "-i", file_path],
            capture_output=True, text=True, timeout=10
        )
        # Duration is in stderr for ffmpeg
        stderr = result.stderr
        import re
        duration_match = re.search(r'Duration:\s*(\d+:\d+:\d+\.\d+)', stderr)
        if duration_match:
            duration_info = f"Duration: {duration_match.group(1)}"
    except Exception:
        pass

    return (
        f"Media file uploaded: {original_filename}\n"
        f"File size: {file_size_mb:.2f} MB\n"
        f"File type: {ext or 'unknown'}\n"
        f"{duration_info}\n\n"
        f"This file has been stored successfully. To enable full AI-powered transcription, "
        f"configure the SARVAM_API_KEY in your backend .env file.\n\n"
        f"The AuraOS chat system can still reference this document by filename and metadata."
    )


# ---------------------------------------------------------------------------
# Embedding pipeline
# ---------------------------------------------------------------------------

def embed_and_store(transcription: str, file_id: str, organization_id: str):
    """Chunk, embed with Gemini, and store in Pinecone. Gracefully skips if unavailable."""
    if not gemini_client:
        print(f"[{file_id}] Skipping embedding: Gemini client not available.")
        return

    if not pinecone_index:
        print(f"[{file_id}] Skipping vector storage: Pinecone not available.")
        return

    text_chunks = chunk_text(transcription)
    vectors_to_upsert = []

    for i, chunk in enumerate(text_chunks):
        if not chunk.strip():
            continue
        try:
            embed_result = gemini_client.models.embed_content(
                model="gemini-embedding-2",
                contents=chunk
            )
            embedding = embed_result.embeddings[0].values

            vectors_to_upsert.append({
                "id": f"{file_id}_chunk_{i}",
                "values": embedding,
                "metadata": {
                    "text": chunk,
                    "file_id": file_id,
                    "organization_id": organization_id,
                }
            })
        except Exception as embed_err:
            print(f"[{file_id}] Embedding error on chunk {i}: {embed_err}")

    if vectors_to_upsert:
        try:
            pinecone_index.upsert(vectors=vectors_to_upsert)
            print(f"[{file_id}] Successfully embedded {len(vectors_to_upsert)} chunks into Pinecone.")
        except Exception as upsert_err:
            print(f"[{file_id}] Pinecone upsert error: {upsert_err}")


# ---------------------------------------------------------------------------
# Background processing task
# ---------------------------------------------------------------------------

def translate_text(text: str, target_language: str) -> str:
    """Translate text using Gemini if target_language is specified and gemini is available."""
    if not target_language or target_language.lower() in ["original", "none", "auto", ""]:
        return text
    if not gemini_client:
        return text + f"\n\n[Translation to {target_language} skipped: Gemini client unavailable]"
    
    prompt = f"Translate the following text into {target_language}. Respond ONLY with the translation, maintaining paragraph structure:\n\n{text}"
    try:
        response = gemini_client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt
        )
        return response.text.strip()
    except Exception as e:
        print(f"Translation to {target_language} failed: {e}")
        return text + f"\n\n[Translation to {target_language} failed: {e}]"


# ---------------------------------------------------------------------------
# Background processing task
# ---------------------------------------------------------------------------

def process_audio(file_path: str, file_id: str, original_filename: str, organization_id: str, target_language: str = "original"):
    """Process uploaded media in the background: transcribe, chunk, embed, store."""
    from ..database import SessionLocal

    db = SessionLocal()
    print(f"Starting processing for {file_id} (org: {organization_id}, language: {target_language})")

    try:
        # Step 1: Transcribe / Extract Text
        ext = os.path.splitext(original_filename)[1].lower()
        is_media = ext in [".mp3", ".wav", ".m4a", ".mp4", ".mpeg", ".ogg", ".aac"]

        if is_media and SARVAM_API_KEY:
            transcription = transcribe_with_sarvam(file_path, file_id)
        else:
            transcription = transcribe_fallback(file_path, file_id, original_filename)

        # Step 1.5: Translate if requested
        if target_language and target_language.lower() != "original":
            transcription = translate_text(transcription, target_language)

        print(f"Processing successful for {file_id}: {transcription[:100]}...")

        # Save JSON for legacy compatibility
        with open(f"uploads/{file_id}.json", "w") as f:
            json.dump({
                "transcription": transcription,
                "filename": original_filename,
                "file_id": file_id
            }, f)

        # Update document record in SQL
        doc = db.query(Document).filter(Document.file_id == file_id).first()
        if doc:
            doc.transcription = transcription
            doc.status = "completed"
            db.commit()

        # Step 2: Embed and store (gracefully skips if services unavailable)
        embed_and_store(transcription, file_id, organization_id)

    except Exception as e:
        print(f"[{file_id}] Processing error: {e}")
        doc = db.query(Document).filter(Document.file_id == file_id).first()
        if doc:
            doc.status = "error"
            doc.error_message = str(e)
            db.commit()

        with open(f"uploads/{file_id}.json", "w") as f:
            json.dump({"error": f"Processing Error: {str(e)}"}, f)
    finally:
        db.close()


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------

@router.post("/", response_model=UploadResponse)
async def upload_audio(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    language: str = Form("original"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Upload media for transcription. Requires authentication."""
    # Check subscription limits
    sub = db.query(Subscription).filter(
        Subscription.organization_id == user.organization_id
    ).first()

    doc_count = db.query(Document).filter(
        Document.organization_id == user.organization_id
    ).count()

    if sub and doc_count >= sub.max_documents:
        raise HTTPException(
            status_code=403,
            detail=f"Document limit reached ({sub.max_documents}). Upgrade your plan."
        )

    file_id = str(uuid.uuid4())
    file_path = os.path.join(UPLOAD_DIR, f"{file_id}_{file.filename}")

    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())

    # Create document record in SQL
    doc = Document(
        id=generate_uuid(),
        file_id=file_id,
        filename=file.filename,
        status="processing",
        organization_id=user.organization_id,
    )
    db.add(doc)
    db.commit()

    background_tasks.add_task(
        process_audio, file_path, file_id, file.filename, user.organization_id, language
    )

    mode = "full" if SARVAM_API_KEY else "local"
    return UploadResponse(
        file_id=file_id,
        filename=file.filename,
        message=f"File uploaded successfully. Processing in background ({mode} mode)."
    )


@router.get("/{file_id}/transcription")
def get_transcription(
    file_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get transcription status/result — scoped to user's org."""
    doc = db.query(Document).filter(
        Document.file_id == file_id,
        Document.organization_id == user.organization_id,
    ).first()

    if doc:
        if doc.status == "error":
            return {"error": doc.error_message or "Processing failed."}
        if doc.status == "completed" and doc.transcription:
            return {
                "transcription": doc.transcription,
                "filename": doc.filename,
                "file_id": doc.file_id,
            }
        return {"transcription": None, "status": doc.status}

    # Fallback to legacy JSON file
    file_path = f"uploads/{file_id}.json"
    if os.path.exists(file_path):
        with open(file_path, "r") as f:
            return json.load(f)
    return {"transcription": None}


@router.get("/files")
def list_files(
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List all uploaded documents for the user's organization."""
    docs = db.query(Document).filter(
        Document.organization_id == user.organization_id,
        Document.status == "completed",
    ).order_by(Document.created_at.desc()).all()

    files_data = []
    for doc in docs:
        files_data.append({
            "file_id": doc.file_id,
            "filename": doc.filename,
            "snippet": (doc.transcription[:100] + "...") if doc.transcription else "",
            "status": doc.status,
            "created_at": doc.created_at.isoformat() if doc.created_at else None,
        })

    # Also include legacy JSON files for backward compat
    if not files_data:
        for filepath in glob.glob("uploads/*.json"):
            try:
                with open(filepath, "r") as f:
                    data = json.load(f)
                    if "transcription" in data and data["transcription"]:
                        files_data.append({
                            "file_id": data.get("file_id", os.path.basename(filepath).replace(".json", "")),
                            "filename": data.get("filename", "Unknown File"),
                            "snippet": data["transcription"][:100] + "...",
                        })
            except Exception:
                pass

    return {"files": files_data}
